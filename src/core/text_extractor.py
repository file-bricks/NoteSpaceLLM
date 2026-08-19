#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Text Extractor - Extract text from various document formats
============================================================

Supports: PDF, DOCX, TXT, MD, XLSX, MSG, EML, and more.
Includes OCR fallback for image-based PDFs.
"""

import email
from email import policy
from email.parser import BytesParser
import html
from html.parser import HTMLParser
from pathlib import Path
import re
from typing import Optional, Tuple
from dataclasses import dataclass


@dataclass
class ExtractionResult:
    """Result of text extraction."""
    success: bool
    text: str
    word_count: int = 0
    error: str = ""
    method: str = ""  # e.g., "native", "ocr", "conversion"


class TextExtractor:
    """
    Extract text content from various document formats.

    Supports:
    - Plain text: .txt, .md, .rst, .log
    - Documents: .pdf, .docx, .doc, .rtf
    - Data: .json, .xml, .yaml, .csv
    - Spreadsheets: .xlsx, .xls
    - Email: .eml, .msg
    - Code: .py, .js, .java, etc.
    """

    def __init__(self, enable_ocr: bool = True):
        """
        Initialize the extractor.

        Args:
            enable_ocr: Enable OCR for image-based PDFs
        """
        self.enable_ocr = enable_ocr
        self._check_dependencies()

    def _check_dependencies(self) -> dict:
        """Check which optional dependencies are available."""
        self._deps = {
            "docx": False,
            "fitz": False,
            "openpyxl": False,
            "pytesseract": False,
            "extract_msg": False,
        }

        try:
            import docx
            self._deps["docx"] = True
        except ImportError:
            pass

        try:
            import fitz
            self._deps["fitz"] = True
        except ImportError:
            pass

        try:
            import openpyxl
            self._deps["openpyxl"] = True
        except ImportError:
            pass

        try:
            import pytesseract
            from PIL import Image
            self._deps["pytesseract"] = True
        except ImportError:
            pass

        try:
            import extract_msg
            self._deps["extract_msg"] = True
        except ImportError:
            pass

        return self._deps

    def extract(self, filepath: Path) -> ExtractionResult:
        """
        Extract text from a file.

        Args:
            filepath: Path to the file

        Returns:
            ExtractionResult with text content or error
        """
        filepath = Path(filepath)

        if not filepath.exists():
            return ExtractionResult(False, "", error=f"File not found: {filepath}")

        if filepath.is_dir():
            return ExtractionResult(False, "", error="Cannot extract text from directory")

        suffix = filepath.suffix.lower()

        try:
            # Plain text files
            if suffix in {".txt", ".md", ".rst", ".log", ".py", ".js", ".ts",
                         ".java", ".cpp", ".c", ".h", ".json", ".xml", ".yaml",
                         ".yml", ".csv", ".html", ".css"}:
                return self._extract_text_file(filepath)

            # PDF
            elif suffix == ".pdf":
                return self._extract_pdf(filepath)

            # Word documents
            elif suffix == ".docx":
                return self._extract_docx(filepath)

            elif suffix == ".doc":
                return self._extract_doc(filepath)

            # RTF (treat as plain text after stripping control codes)
            elif suffix == ".rtf":
                return self._extract_rtf(filepath)

            # PowerPoint
            elif suffix == ".pptx":
                return self._extract_pptx(filepath)

            # Excel
            elif suffix in {".xlsx", ".xls"}:
                return self._extract_excel(filepath)

            # Email
            elif suffix == ".eml":
                return self._extract_eml(filepath)

            elif suffix == ".msg":
                return self._extract_msg(filepath)

            else:
                return ExtractionResult(
                    False, "",
                    error=f"Unsupported file format: {suffix}"
                )

        except Exception as e:
            return ExtractionResult(False, "", error=str(e))

    def _extract_text_file(self, filepath: Path) -> ExtractionResult:
        """Extract text from plain text files."""
        try:
            text = filepath.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                text = filepath.read_text(encoding="latin-1")
            except Exception as e:
                return ExtractionResult(False, "", error=f"Encoding error: {e}")

        return ExtractionResult(
            success=True,
            text=text,
            word_count=len(text.split()),
            method="native"
        )

    def _extract_pdf(self, filepath: Path) -> ExtractionResult:
        """Extract text from PDF, with OCR fallback."""
        if not self._deps["fitz"]:
            return ExtractionResult(
                False, "",
                error="PyMuPDF (fitz) not installed. Install with: pip install pymupdf"
            )

        import fitz

        try:
            doc = fitz.open(str(filepath))
            text_parts = []
            pages_with_text = 0

            for page in doc:
                page_text = page.get_text().strip()
                if page_text:
                    text_parts.append(page_text)
                    pages_with_text += 1

            doc.close()

            # If no text found, try OCR
            if not text_parts and self.enable_ocr:
                return self._extract_pdf_ocr(filepath)

            text = "\n\n".join(text_parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )

        except Exception as e:
            return ExtractionResult(False, "", error=f"PDF error: {e}")

    def _extract_pdf_ocr(self, filepath: Path) -> ExtractionResult:
        """Extract text from image-based PDF using OCR."""
        if not self._deps["pytesseract"]:
            return ExtractionResult(
                False, "",
                error="OCR not available. Install pytesseract and Tesseract."
            )

        if not self._deps["fitz"]:
            return ExtractionResult(False, "", error="PyMuPDF required for OCR")

        import fitz
        import pytesseract
        from PIL import Image
        import io

        try:
            doc = fitz.open(str(filepath))
            text_parts = []

            for page_num, page in enumerate(doc):
                # Render page to image
                mat = fitz.Matrix(2, 2)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")

                # OCR the image
                image = Image.open(io.BytesIO(img_data))
                page_text = pytesseract.image_to_string(image, lang="deu+eng")

                if page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

            doc.close()

            if not text_parts:
                return ExtractionResult(False, "", error="OCR found no text")

            text = "\n\n".join(text_parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="ocr"
            )

        except Exception as e:
            return ExtractionResult(False, "", error=f"OCR error: {e}")

    def _extract_docx(self, filepath: Path) -> ExtractionResult:
        """Extract text from Word document (.docx)."""
        if not self._deps["docx"]:
            return ExtractionResult(
                False, "",
                error="python-docx not installed. Install with: pip install python-docx"
            )

        from docx import Document

        try:
            doc = Document(str(filepath))
            parts = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    # row.cells liefert eine ueber mehrere Rasterspalten verbundene
                    # Zelle einmal PRO SPALTE (python-docx-Verhalten) -- ohne
                    # Dedup ueber die Identitaet des zugrunde liegenden <w:tc>-
                    # Elements wird der Zellinhalt so oft wiederholt, wie er
                    # Spalten ueberspannt (Nebenfund aus T-20260817-825816579,
                    # dort in BACH DocumentPipeline._extract_docx behoben:
                    # Faktor 17,8x an einer realen Formular-DOCX gemessen).
                    seen: set = set()
                    cells = []
                    for c in row.cells:
                        key = id(c._tc)
                        if key in seen:
                            continue
                        seen.add(key)
                        if c.text.strip():
                            cells.append(c.text.strip())
                    if cells:
                        parts.append(" | ".join(cells))

            text = "\n\n".join(parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )

        except Exception as e:
            return ExtractionResult(False, "", error=f"DOCX error: {e}")

    def _extract_doc(self, filepath: Path) -> ExtractionResult:
        """Extract text from old Word format (.doc)."""
        import subprocess
        import shutil

        # Try antiword first
        if shutil.which("antiword"):
            try:
                result = subprocess.run(
                    ["antiword", str(filepath)],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    text = result.stdout.strip()
                    return ExtractionResult(
                        success=True,
                        text=text,
                        word_count=len(text.split()),
                        method="antiword"
                    )
            except Exception:
                pass

        # Try LibreOffice conversion
        if shutil.which("soffice"):
            try:
                import tempfile
                with tempfile.TemporaryDirectory() as tmpdir:
                    subprocess.run(
                        ["soffice", "--headless", "--convert-to", "txt:Text",
                         "--outdir", tmpdir, str(filepath)],
                        capture_output=True,
                        timeout=60
                    )
                    txt_file = Path(tmpdir) / (filepath.stem + ".txt")
                    if txt_file.exists():
                        text = txt_file.read_text(encoding="utf-8", errors="replace")
                        return ExtractionResult(
                            success=True,
                            text=text,
                            word_count=len(text.split()),
                            method="libreoffice"
                        )
            except Exception:
                pass

        return ExtractionResult(
            False, "",
            error=".doc extraction requires antiword or LibreOffice"
        )

    @staticmethod
    def _html_to_text(html_content: str) -> str:
        """Strip HTML tags and convert entities to plain text."""
        if not html_content or not html_content.strip():
            return ""

        class _HTMLParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.result = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                if tag in ('script', 'style', 'head', 'title'):
                    self._skip = True
                elif tag in ('p', 'br', 'div', 'tr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'):
                    self.result.append('\n')

            def handle_endtag(self, tag):
                if tag in ('script', 'style', 'head', 'title'):
                    self._skip = False
                elif tag in ('p', 'div', 'tr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'):
                    self.result.append('\n')

            def handle_data(self, data):
                if not self._skip:
                    self.result.append(data)

            def get_text(self) -> str:
                text = ''.join(self.result)
                lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
                return '\n'.join(line for line in lines if line)

        try:
            parser = _HTMLParser()
            parser.feed(html_content)
            parsed_text = parser.get_text()
            if parsed_text.strip():
                return parsed_text
        except Exception:
            pass

        # Fallback regex strip
        clean = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r'<[^>]+>', ' ', clean)
        clean = html.unescape(clean)
        lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in clean.split('\n')]
        return '\n'.join(line for line in lines if line)

    @staticmethod
    def _parse_rtf_bytes(raw_bytes: bytes) -> str:
        """Extract plain text from RTF byte stream, handling codepages, hex escapes and unicode."""
        text = raw_bytes.decode('latin-1', errors='replace')

        cp_match = re.search(r'\\ansicpg(\d+)', text)
        encoding = f'cp{cp_match.group(1)}' if cp_match else 'cp1252'

        skip_destinations = {
            'fonttbl', 'colortbl', 'stylesheet', 'info', 'generator',
            'pict', 'header', 'footer', 'headerl', 'headerr', 'headerf',
            'footerl', 'footerr', 'footerf', 'object', 'template', 'themedata'
        }

        stack = []
        ignorable = False
        skip_depth = 0
        uc_skip = 1

        out = []
        i = 0
        n = len(text)

        while i < n:
            c = text[i]
            if c == '{':
                stack.append((skip_depth, uc_skip))
                i += 1
            elif c == '}':
                if stack:
                    skip_depth, uc_skip = stack.pop()
                i += 1
            elif c == '\\':
                i += 1
                if i >= n:
                    break
                ch = text[i]
                if ch in ('\\', '{', '}'):
                    if skip_depth == 0:
                        out.append(ch)
                    i += 1
                elif ch == '~':
                    if skip_depth == 0:
                        out.append(' ')
                    i += 1
                elif ch == '_':
                    if skip_depth == 0:
                        out.append('-')
                    i += 1
                elif ch == '*':
                    i += 1
                    ignorable = True
                elif ch == "'":
                    hex_str = text[i+1:i+3]
                    i += 3
                    if skip_depth == 0 and len(hex_str) == 2:
                        try:
                            byte_val = bytes.fromhex(hex_str)
                            out.append(byte_val.decode(encoding, errors='replace'))
                        except Exception:
                            pass
                else:
                    match = re.match(r'([a-zA-Z]+)(-?\d+)? ?', text[i:])
                    if match:
                        word = match.group(1)
                        param = match.group(2)
                        full_len = match.end()
                        i += full_len

                        if word in ('par', 'line', 'row', 'sect'):
                            if skip_depth == 0:
                                out.append('\n')
                        elif word in ('tab', 'cell'):
                            if skip_depth == 0:
                                out.append('\t')
                        elif word == 'uc':
                            if param:
                                uc_skip = max(0, int(param))
                        elif word == 'u':
                            if param:
                                code_point = int(param)
                                if code_point < 0:
                                    code_point += 65536
                                if skip_depth == 0:
                                    try:
                                        out.append(chr(code_point))
                                    except Exception:
                                        pass
                                i += uc_skip
                        elif word in skip_destinations or ignorable:
                            skip_depth += 1
                        ignorable = False
                    else:
                        i += 1
            else:
                if skip_depth == 0 and c not in ('\r', '\n'):
                    out.append(c)
                i += 1

        res = ''.join(out)
        lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in res.split('\n')]
        return '\n'.join(line for line in lines if line)

    def _extract_rtf(self, filepath: Path) -> ExtractionResult:
        """Extract text from RTF files with proper control code stripping, unicode and hex escape decoding."""
        try:
            raw = filepath.read_bytes()
            text = self._parse_rtf_bytes(raw)

            if not text:
                return ExtractionResult(False, "", error="RTF enthielt keinen extrahierbaren Text")

            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )
        except Exception as e:
            return ExtractionResult(False, "", error=f"RTF error: {e}")

    def _extract_pptx(self, filepath: Path) -> ExtractionResult:
        """Extract text from PowerPoint (.pptx) files."""
        try:
            from zipfile import ZipFile
            import xml.etree.ElementTree as ET

            parts = []
            with ZipFile(str(filepath)) as z:
                # pptx slides are in ppt/slides/slide*.xml
                slide_names = sorted([
                    n for n in z.namelist()
                    if n.startswith("ppt/slides/slide") and n.endswith(".xml")
                ])

                ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}

                for i, slide_name in enumerate(slide_names, 1):
                    with z.open(slide_name) as f:
                        tree = ET.parse(f)
                    texts = [t.text for t in tree.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}t") if t.text]
                    if texts:
                        parts.append(f"--- Slide {i} ---")
                        parts.append("\n".join(texts))

            if not parts:
                return ExtractionResult(False, "", error="Keine Textinhalte in PPTX gefunden")

            text = "\n\n".join(parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )
        except Exception as e:
            return ExtractionResult(False, "", error=f"PPTX error: {e}")

    def _extract_excel(self, filepath: Path) -> ExtractionResult:
        """Extract text from Excel files."""
        if not self._deps["openpyxl"]:
            return ExtractionResult(
                False, "",
                error="openpyxl not installed. Install with: pip install openpyxl"
            )

        import openpyxl

        try:
            wb = openpyxl.load_workbook(str(filepath), data_only=True)
            parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                parts.append(f"[Sheet: {sheet_name}]")

                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))

            wb.close()

            text = "\n".join(parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )

        except Exception as e:
            return ExtractionResult(False, "", error=f"Excel error: {e}")

    def _extract_eml(self, filepath: Path) -> ExtractionResult:
        """Extract text from .eml email files."""
        try:
            with open(filepath, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)

            parts = []

            if msg['Date']:
                parts.append(f"Date: {msg['Date']}")
            if msg['From']:
                parts.append(f"From: {msg['From']}")
            if msg['To']:
                parts.append(f"To: {msg['To']}")
            if msg['Subject']:
                parts.append(f"Subject: {msg['Subject']}")

            parts.append("")

            def _decode_part(part) -> str:
                charset = part.get_content_charset() or 'utf-8'
                payload = part.get_payload(decode=True)
                if isinstance(payload, bytes):
                    for enc in (charset, 'utf-8', 'latin-1', 'cp1252'):
                        try:
                            return payload.decode(enc)
                        except (UnicodeDecodeError, LookupError):
                            continue
                    return payload.decode('utf-8', errors='replace')
                return str(payload or '')

            # Extract body (plain text prioritized, HTML fallback)
            plain_body = ""
            html_body = ""

            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/plain' and not plain_body:
                        plain_body = _decode_part(part)
                    elif content_type == 'text/html' and not html_body:
                        html_body = _decode_part(part)
            else:
                content_type = msg.get_content_type()
                if content_type == 'text/html':
                    html_body = _decode_part(msg)
                else:
                    plain_body = _decode_part(msg)

            if plain_body.strip():
                parts.append(plain_body.strip())
            elif html_body.strip():
                extracted_html_text = self._html_to_text(html_body)
                if extracted_html_text:
                    parts.append(extracted_html_text)

            text = "\n".join(parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )

        except Exception as e:
            return ExtractionResult(False, "", error=f"EML error: {e}")

    def _extract_msg(self, filepath: Path) -> ExtractionResult:
        """Extract text from Outlook .msg files."""
        if not self._deps["extract_msg"]:
            return ExtractionResult(
                False, "",
                error="extract-msg not installed. Install with: pip install extract-msg"
            )

        import extract_msg

        try:
            msg = extract_msg.Message(str(filepath))
            parts = []

            if msg.date:
                parts.append(f"Date: {msg.date}")
            if msg.sender:
                parts.append(f"From: {msg.sender}")
            if msg.to:
                parts.append(f"To: {msg.to}")
            if msg.subject:
                parts.append(f"Subject: {msg.subject}")

            parts.append("")

            body_text = ""
            if msg.body:
                body_text = msg.body
            elif getattr(msg, 'htmlBody', None):
                html_raw = msg.htmlBody
                if isinstance(html_raw, bytes):
                    html_raw = html_raw.decode('utf-8', errors='replace')
                body_text = self._html_to_text(str(html_raw))

            if body_text:
                parts.append(body_text)

            msg.close()

            text = "\n".join(parts)
            return ExtractionResult(
                success=True,
                text=text,
                word_count=len(text.split()),
                method="native"
            )

        except Exception as e:
            return ExtractionResult(False, "", error=f"MSG error: {e}")

    def get_dependencies_status(self) -> dict:
        """Get status of optional dependencies."""
        return {
            "python-docx": {"installed": self._deps["docx"], "for": "DOCX files"},
            "PyMuPDF": {"installed": self._deps["fitz"], "for": "PDF files"},
            "openpyxl": {"installed": self._deps["openpyxl"], "for": "Excel files"},
            "pytesseract": {"installed": self._deps["pytesseract"], "for": "OCR (image PDFs)"},
            "extract-msg": {"installed": self._deps["extract_msg"], "for": "Outlook MSG files"},
        }
