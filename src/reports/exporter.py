#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Report Exporter - Export reports to various formats
===================================================

Supports: Markdown, PDF, DOCX, HTML, TXT
"""

import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass


@dataclass
class ExportResult:
    """Result of an export operation."""
    success: bool
    filepath: Path = None
    format: str = ""
    error: str = ""


class ReportExporter:
    """
    Export reports to multiple formats.

    Supported formats:
    - md: Markdown (native)
    - txt: Plain text
    - html: HTML (simple conversion)
    - pdf: PDF (requires pandoc or weasyprint)
    - docx: Word (requires python-docx)
    """

    def __init__(self, output_directory: Path):
        """
        Initialize the exporter.

        Args:
            output_directory: Base directory for exports
        """
        self.output_dir = Path(output_directory)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        content: str,
        base_name: str,
        formats: List[str],
        title: str = "",
        author: str = ""
    ) -> List[ExportResult]:
        """
        Export content to multiple formats.

        Args:
            content: Markdown content to export
            base_name: Base filename (without extension)
            formats: List of formats to export to
            title: Optional document title
            author: Optional author name

        Returns:
            List of ExportResult for each format
        """
        results = []
        safe_name = self._safe_filename(base_name)

        for fmt in formats:
            fmt = fmt.lower().strip(".")

            if fmt == "md":
                result = self._export_markdown(content, safe_name, title, author)
            elif fmt == "txt":
                result = self._export_text(content, safe_name)
            elif fmt == "html":
                result = self._export_html(content, safe_name, title)
            elif fmt == "pdf":
                result = self._export_pdf(content, safe_name, title, author)
            elif fmt == "docx":
                result = self._export_docx(content, safe_name, title, author)
            else:
                result = ExportResult(False, format=fmt, error=f"Unbekanntes Format: {fmt}")

            results.append(result)

        return results

    def _safe_filename(self, name: str) -> str:
        """Create a safe filename."""
        safe = "".join(c for c in name if c.isalnum() or c in " -_").strip()
        safe = safe.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{safe}_{timestamp}" if safe else f"report_{timestamp}"

    def _export_markdown(self, content: str, name: str, title: str, author: str) -> ExportResult:
        """Export to Markdown."""
        try:
            filepath = self.output_dir / f"{name}.md"

            # Add metadata header
            header = ""
            if title or author:
                header = "---\n"
                if title:
                    header += f"title: {title}\n"
                if author:
                    header += f"author: {author}\n"
                header += f"date: {datetime.now().strftime('%Y-%m-%d')}\n"
                header += "---\n\n"

            filepath.write_text(header + content, encoding="utf-8")

            return ExportResult(True, filepath, "md")

        except Exception as e:
            return ExportResult(False, format="md", error=str(e))

    def _export_text(self, content: str, name: str) -> ExportResult:
        """Export to plain text."""
        try:
            filepath = self.output_dir / f"{name}.txt"

            # Strip Markdown formatting
            plain = content
            plain = re.sub(r'^#+\s+', '', plain, flags=re.MULTILINE)  # Headers
            plain = re.sub(r'\*\*(.+?)\*\*', r'\1', plain)  # Bold
            plain = re.sub(r'\*(.+?)\*', r'\1', plain)  # Italic
            plain = re.sub(r'`(.+?)`', r'\1', plain)  # Code
            plain = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', plain)  # Links
            plain = re.sub(r'^\s*[-*]\s+', '- ', plain, flags=re.MULTILINE)  # Lists

            filepath.write_text(plain, encoding="utf-8")

            return ExportResult(True, filepath, "txt")

        except Exception as e:
            return ExportResult(False, format="txt", error=str(e))

    def _export_html(self, content: str, name: str, title: str) -> ExportResult:
        """Export to HTML."""
        try:
            filepath = self.output_dir / f"{name}.html"

            # Convert Markdown to HTML
            html_content = self._markdown_to_html(content)

            html = f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{title or 'Bericht'}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{ border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #2c3e50; margin-top: 30px; }}
        h3 {{ color: #34495e; }}
        code {{
            background: #f5f5f5;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Consolas', monospace;
        }}
        pre {{
            background: #f5f5f5;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 0;
            padding-left: 20px;
            color: #666;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }}
        th {{ background: #f5f5f5; }}
    </style>
</head>
<body>
{html_content}
<hr>
<p><small>Generiert am {datetime.now().strftime('%d.%m.%Y %H:%M')} mit NoteSpaceLLM</small></p>
</body>
</html>"""

            filepath.write_text(html, encoding="utf-8")

            return ExportResult(True, filepath, "html")

        except Exception as e:
            return ExportResult(False, format="html", error=str(e))

    def _markdown_to_html(self, markdown: str) -> str:
        """Convert Markdown to HTML."""
        html = markdown

        # Headers
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

        # Bold and Italic
        html = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', html)
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

        # Code
        html = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code>\2</code></pre>', html, flags=re.DOTALL)
        html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)

        # Links
        html = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', html)

        # Lists
        lines = html.split('\n')
        in_list = False
        result = []
        for line in lines:
            if re.match(r'^\s*[-*]\s+', line):
                if not in_list:
                    result.append('<ul>')
                    in_list = True
                line = re.sub(r'^\s*[-*]\s+', '', line)
                result.append(f'<li>{line}</li>')
            else:
                if in_list:
                    result.append('</ul>')
                    in_list = False
                result.append(line)

        if in_list:
            result.append('</ul>')

        html = '\n'.join(result)

        # Paragraphs
        html = re.sub(r'\n\n+', r'</p>\n<p>', html)
        html = f'<p>{html}</p>'

        # Clean up
        html = re.sub(r'<p>\s*(<h[1-6]>)', r'\1', html)
        html = re.sub(r'(</h[1-6]>)\s*</p>', r'\1', html)
        html = re.sub(r'<p>\s*(<ul>)', r'\1', html)
        html = re.sub(r'(</ul>)\s*</p>', r'\1', html)
        html = re.sub(r'<p>\s*</p>', '', html)

        return html

    def _export_pdf(self, content: str, name: str, title: str, author: str) -> ExportResult:
        """Export to PDF using pandoc or weasyprint."""
        try:
            filepath = self.output_dir / f"{name}.pdf"

            # Try pandoc first
            if self._try_pandoc_pdf(content, filepath, title, author):
                return ExportResult(True, filepath, "pdf")

            # Try weasyprint
            if self._try_weasyprint_pdf(content, filepath, title):
                return ExportResult(True, filepath, "pdf")

            return ExportResult(
                False, format="pdf",
                error="PDF-Export benoetigt pandoc oder weasyprint"
            )

        except Exception as e:
            return ExportResult(False, format="pdf", error=str(e))

    def _try_pandoc_pdf(self, content: str, filepath: Path, title: str, author: str) -> bool:
        """Try to export PDF with pandoc."""
        try:
            import shutil
            if not shutil.which("pandoc"):
                return False

            with tempfile.NamedTemporaryFile(mode='w', suffix='.md',
                                            delete=False, encoding='utf-8') as f:
                # Add metadata
                header = "---\n"
                if title:
                    header += f"title: '{title}'\n"
                if author:
                    header += f"author: '{author}'\n"
                header += f"date: '{datetime.now().strftime('%d.%m.%Y')}'\n"
                header += "---\n\n"

                f.write(header + content)
                md_path = f.name

            result = subprocess.run(
                ['pandoc', md_path, '-o', str(filepath), '--pdf-engine=xelatex'],
                capture_output=True,
                timeout=60
            )

            Path(md_path).unlink(missing_ok=True)
            return result.returncode == 0 and filepath.exists()

        except Exception:
            return False

    def _try_weasyprint_pdf(self, content: str, filepath: Path, title: str) -> bool:
        """Try to export PDF with weasyprint."""
        try:
            from weasyprint import HTML

            html_content = self._markdown_to_html(content)
            html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title or 'Bericht'}</title>
    <style>
        @page {{ margin: 2cm; }}
        body {{ font-family: sans-serif; line-height: 1.6; }}
        h1 {{ color: #2c3e50; }}
    </style>
</head>
<body>{html_content}</body>
</html>"""

            HTML(string=html).write_pdf(str(filepath))
            return True

        except ImportError:
            return False
        except Exception:
            return False

    def _export_docx(self, content: str, name: str, title: str, author: str) -> ExportResult:
        """Export to Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            filepath = self.output_dir / f"{name}.docx"

            doc = Document()

            # Title
            if title:
                title_para = doc.add_heading(title, level=0)
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Process content
            for line in content.split('\n'):
                line = line.strip()

                if not line:
                    continue

                # Headers
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                # List items
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                # Regular paragraphs
                else:
                    # Clean up markdown
                    clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                    clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                    clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)
                    doc.add_paragraph(clean_line)

            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.add_run(f"Generiert am {datetime.now().strftime('%d.%m.%Y %H:%M')} mit NoteSpaceLLM")
            footer.runs[0].font.size = Pt(8)
            footer.runs[0].font.italic = True

            doc.save(str(filepath))

            return ExportResult(True, filepath, "docx")

        except ImportError:
            return ExportResult(
                False, format="docx",
                error="python-docx nicht installiert. Installieren mit: pip install python-docx"
            )
        except Exception as e:
            return ExportResult(False, format="docx", error=str(e))
