#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Main Window - Central application window for NoteSpaceLLM
=========================================================

Integrates all panels and manages the application flow.
Includes RAG integration with LangChain + ChromaDB.
"""

import sys
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

try:
    from PySide6.QtWidgets import (
        QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QSplitter,
        QMenuBar, QMenu, QToolBar, QStatusBar, QMessageBox,
        QFileDialog, QInputDialog, QApplication, QProgressDialog,
        QDockWidget, QTabWidget, QDialog, QComboBox, QLabel,
        QDialogButtonBox, QFormLayout, QGroupBox, QLineEdit,
        QPushButton
    )
    from PySide6.QtCore import Qt, QSize, QTimer, Slot
    from PySide6.QtGui import QAction, QIcon, QKeySequence
    PYSIDE_AVAILABLE = True
except ImportError:
    PYSIDE_AVAILABLE = False


if PYSIDE_AVAILABLE:
    from PySide6.QtCore import QThread, Signal

    class AnalysisWorker(QThread):
        """Worker thread for batch sub-query analysis."""
        query_complete = Signal(str, str, str)  # query_id, response, error
        all_complete = Signal()

        def __init__(self, llm_client, tasks):
            super().__init__()
            self._llm_client = llm_client
            self._tasks = tasks  # [(query_id, prompt), ...]

        def run(self):
            for query_id, prompt in self._tasks:
                try:
                    response = self._llm_client.chat(prompt, "")
                    self.query_complete.emit(query_id, response, "")
                except Exception as e:
                    self.query_complete.emit(query_id, "", str(e))
            self.all_complete.emit()

    class ExtractionWorker(QThread):
        """Worker thread for text extraction (prevents GUI freeze)."""
        progress = Signal(int, int, str)  # current, total, filename
        doc_extracted = Signal(str, str, str)  # doc_id, text, error
        all_complete = Signal()

        def __init__(self, extractor, docs):
            super().__init__()
            self._extractor = extractor
            self._docs = docs  # [(doc_id, doc_path, doc_name), ...]

        def run(self):
            total = len(self._docs)
            for i, (doc_id, doc_path, doc_name) in enumerate(self._docs):
                self.progress.emit(i, total, doc_name)
                try:
                    result = self._extractor.extract(doc_path)
                    if result.success:
                        self.doc_extracted.emit(doc_id, result.text, "")
                    else:
                        self.doc_extracted.emit(doc_id, "", result.error or "Extraktion fehlgeschlagen")
                except Exception as e:
                    self.doc_extracted.emit(doc_id, "", str(e))
            self.progress.emit(total, total, "")
            self.all_complete.emit()

    class IndexWorker(QThread):
        """Worker thread for RAG indexing (prevents GUI freeze)."""
        progress = Signal(int, int, str)  # current, total, doc_name
        doc_indexed = Signal(str, bool)  # doc_id, success
        all_complete = Signal(int, int)  # indexed_count, total_count

        def __init__(self, doc_manager, doc_ids_and_names):
            super().__init__()
            self._doc_manager = doc_manager
            self._docs = doc_ids_and_names  # [(doc_id, doc_name), ...]

        def run(self):
            total = len(self._docs)
            indexed = 0
            for i, (doc_id, doc_name) in enumerate(self._docs):
                self.progress.emit(i, total, doc_name)
                try:
                    success = self._doc_manager.index_document(doc_id)
                    self.doc_indexed.emit(doc_id, success)
                    if success:
                        indexed += 1
                except Exception:
                    self.doc_indexed.emit(doc_id, False)
            self.progress.emit(total, total, "")
            self.all_complete.emit(indexed, total)

    class ModelLoadWorker(QThread):
        """Worker thread for Ollama model discovery (prevents GUI freeze on network call)."""
        models_loaded = Signal(list, str)  # models, status_message

        def __init__(self, url: str, api_key: str = ""):
            super().__init__()
            self._url = url
            self._api_key = api_key

        def run(self):
            try:
                from ..llm.ollama_client import OllamaClient
                client = OllamaClient.__new__(OllamaClient)
                client.base_url = self._url.rstrip("/")
                client.api_key = self._api_key
                client._is_available = False
                client.model = ""
                client._check_availability()
                if client._is_available:
                    models = client.get_models()
                    self.models_loaded.emit(models, f"{len(models)} Modelle auf {self._url}")
                else:
                    self.models_loaded.emit([], f"Ollama nicht erreichbar ({self._url})")
            except Exception as e:
                self.models_loaded.emit([], f"Fehler: {e}")


if PYSIDE_AVAILABLE:
    from PySide6.QtWidgets import QScrollArea, QFrame

    class _PipelineHelpDialog(QDialog):
        """Visueller Pipeline-Hilfe-Dialog — erklärt alle Analyse-Schritte."""

        STEPS = [
            {
                "number": "1",
                "title": "Dokumente hinzufügen",
                "icon": "📁",
                "color": "#3498db",
                "desc": (
                    "<b>+ Dateien</b> oder <b>+ Ordner</b> in der Toolbar klicken.\n\n"
                    "Unterstützte Formate: PDF, DOCX, TXT, MD, XLSX, PPTX, RTF, "
                    "CSV, JSON, XML, EML …\n\n"
                    "Die Textextraktion startet automatisch im Hintergrund."
                ),
                "optional": False,
            },
            {
                "number": "2",
                "title": "Text extrahieren",
                "icon": "🔤",
                "color": "#9b59b6",
                "desc": (
                    "Läuft <b>automatisch</b> nach dem Hinzufügen.\n\n"
                    "Alternativ: <b>Text extrahieren</b> in der Toolbar, "
                    "um alle ausgewählten Dokumente manuell zu verarbeiten.\n\n"
                    "Der Fortschritt wird in der Statusleiste angezeigt."
                ),
                "optional": False,
            },
            {
                "number": "3",
                "title": "Detailanalysen (optional)",
                "icon": "🔍",
                "color": "#f39c12",
                "desc": (
                    "Rechtsklick auf ein Dokument → Sub-Query hinzufügen.\n\n"
                    "Dann <b>Analysieren</b> in der Toolbar drücken.\n\n"
                    "Das LLM beantwortet die Sub-Queries pro Dokument "
                    "und die Ergebnisse fließen in den Bericht ein.\n\n"
                    "<i>Optional — der Bericht funktioniert auch ohne diesen Schritt.</i>"
                ),
                "optional": True,
            },
            {
                "number": "4",
                "title": "Bericht erstellen",
                "icon": "📝",
                "color": "#27ae60",
                "desc": (
                    "<b>Bericht erstellen</b> in der Toolbar klicken.\n\n"
                    "Das LLM kombiniert alle Dokumentinhalte und "
                    "Detailanalysen zu einem strukturierten Bericht.\n\n"
                    "Fehlende Textextraktion wird <b>automatisch</b> "
                    "ausgelöst.\n\n"
                    "Der Bericht erscheint im Ausgabe-Panel rechts."
                ),
                "optional": False,
            },
            {
                "number": "5",
                "title": "Exportieren",
                "icon": "💾",
                "color": "#e74c3c",
                "desc": (
                    "Im <b>Ausgabe-Panel</b> rechts: Format wählen "
                    "(Markdown, PDF, DOCX …) und <b>Exportieren</b>.\n\n"
                    "<b>LLM-Prompt exportieren</b>: Exportiert den an das LLM "
                    "gesendeten Prompt als .md Datei — nützlich für andere LLMs.\n\n"
                    "<b>Workspace exportieren</b> (Datei-Menü): Speichert Projekt, "
                    "Dokument-Metadaten, Auszüge und Bericht als JSON "
                    "für Web/Companion-Apps."
                ),
                "optional": False,
            },
        ]

        def __init__(self, parent=None):
            super().__init__(parent)
            self.setWindowTitle("Pipeline-Übersicht")
            self.setMinimumSize(560, 560)
            self.resize(620, 640)
            self._setup_ui()

        def _setup_ui(self):
            layout = QVBoxLayout(self)
            layout.setSpacing(0)
            layout.setContentsMargins(0, 0, 0, 0)

            # Header
            header = QWidget()
            header.setStyleSheet("background: #1a1a2e; padding: 16px 20px 14px;")
            hlay = QVBoxLayout(header)
            hlay.setSpacing(4)
            title = QLabel("Analyse-Pipeline")
            title.setStyleSheet("font-size: 18px; font-weight: 800; color: #fff; letter-spacing: -0.5px;")
            subtitle = QLabel("So funktioniert NoteSpaceLLM in 5 Schritten")
            subtitle.setStyleSheet("font-size: 12px; color: #7f8c9a;")
            hlay.addWidget(title)
            hlay.addWidget(subtitle)
            layout.addWidget(header)

            # Scrollbarer Inhalt
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setFrameShape(QFrame.Shape.NoFrame)
            scroll.setStyleSheet("QScrollArea { background: #f8f9fa; border: none; }")

            content = QWidget()
            content.setStyleSheet("background: #f8f9fa;")
            clay = QVBoxLayout(content)
            clay.setContentsMargins(16, 16, 16, 16)
            clay.setSpacing(8)

            for i, step in enumerate(self.STEPS):
                card = self._make_step_card(step, is_last=(i == len(self.STEPS) - 1))
                clay.addWidget(card)

            clay.addStretch()
            scroll.setWidget(content)
            layout.addWidget(scroll, stretch=1)

            # Footer
            footer = QWidget()
            footer.setStyleSheet("background: #fff; border-top: 1px solid #e0e0e0; padding: 10px 16px;")
            flay = QHBoxLayout(footer)
            tip = QLabel("💡 Tipp: Bewege den Mauszeiger über die Toolbar-Buttons für Kurzerklärungen.")
            tip.setStyleSheet("font-size: 11px; color: #7f8c9a;")
            close_btn = QPushButton("Schließen")
            close_btn.setStyleSheet(
                "QPushButton { padding: 6px 18px; background: #2c3e50; color: #fff; "
                "border: none; border-radius: 4px; font-weight: 600; font-size: 12px; }"
                "QPushButton:hover { background: #34495e; }"
            )
            close_btn.clicked.connect(self.accept)
            flay.addWidget(tip, stretch=1)
            flay.addWidget(close_btn)
            layout.addWidget(footer)

        def _make_step_card(self, step: dict, is_last: bool) -> QWidget:
            card = QFrame()
            card.setFrameShape(QFrame.Shape.NoFrame)
            card.setStyleSheet(
                "QFrame { background: #fff; border: 1px solid #e8ecf0; "
                "border-radius: 8px; }"
            )

            row = QHBoxLayout(card)
            row.setContentsMargins(14, 12, 14, 12)
            row.setSpacing(12)

            # Step indicator (number circle + vertical connector)
            indicator = QWidget()
            indicator.setFixedWidth(36)
            ilay = QVBoxLayout(indicator)
            ilay.setContentsMargins(0, 0, 0, 0)
            ilay.setSpacing(2)

            circle = QLabel(step["number"])
            circle.setFixedSize(32, 32)
            circle.setAlignment(Qt.AlignmentFlag.AlignCenter)
            circle.setStyleSheet(
                f"background: {step['color']}; color: #fff; "
                "border-radius: 16px; font-weight: 800; font-size: 13px;"
            )
            ilay.addWidget(circle, alignment=Qt.AlignmentFlag.AlignHCenter)

            if not is_last:
                connector = QFrame()
                connector.setFrameShape(QFrame.Shape.VLine)
                connector.setStyleSheet(f"color: {step['color']}; opacity: 0.3;")
                connector.setFixedWidth(2)
                connector.setFixedHeight(12)
                ilay.addWidget(connector, alignment=Qt.AlignmentFlag.AlignHCenter)

            row.addWidget(indicator, alignment=Qt.AlignmentFlag.AlignTop)

            # Content
            content = QWidget()
            clay = QVBoxLayout(content)
            clay.setContentsMargins(0, 0, 0, 0)
            clay.setSpacing(4)

            # Title row
            title_row = QHBoxLayout()
            icon_label = QLabel(f"{step['icon']} {step['title']}")
            icon_label.setStyleSheet(
                f"font-weight: 700; font-size: 13px; color: #1a1a2e;"
            )
            title_row.addWidget(icon_label)
            if step["optional"]:
                badge = QLabel("optional")
                badge.setStyleSheet(
                    "background: #fff3cd; color: #856404; padding: 1px 7px; "
                    "border-radius: 10px; font-size: 10px; font-weight: 600;"
                )
                title_row.addWidget(badge)
            title_row.addStretch()
            clay.addLayout(title_row)

            # Description
            desc = QLabel(step["desc"])
            desc.setWordWrap(True)
            desc.setTextFormat(Qt.TextFormat.RichText)
            desc.setStyleSheet("font-size: 12px; color: #4a5568; line-height: 1.5;")
            clay.addWidget(desc)

            row.addWidget(content, stretch=1)
            return card


class MainWindow(QMainWindow if PYSIDE_AVAILABLE else object):
    """
    Main application window.

    Layout:
    +--------------------------------------------------+
    |  Menu Bar                                        |
    +--------------------------------------------------+
    |  Toolbar                                         |
    +--------------------------------------------------+
    |          |                     |                 |
    | Documents|     Workflow        |     Output      |
    |  Panel   |      Panel          |     Panel       |
    |          |                     |                 |
    |          +---------------------+-----------------+
    |          |                     |                 |
    |          |        Chat         |    Settings     |
    |          |        Panel        |                 |
    |          |                     |                 |
    +--------------------------------------------------+
    |  Status Bar                                      |
    +--------------------------------------------------+
    """

    def __init__(self):
        if not PYSIDE_AVAILABLE:
            raise ImportError("PySide6 is required. Install with: pip install PySide6")

        super().__init__()

        # Initialize components
        from ..core.document_manager import DocumentManager
        from ..core.sub_query import SubQueryManager
        from ..core.project import Project, ProjectManager
        from ..core.text_extractor import TextExtractor

        # Project management
        self._projects_dir = Path.home() / "NoteSpaceLLM" / "projects"
        self._projects_dir.mkdir(parents=True, exist_ok=True)

        self._project_manager = ProjectManager(self._projects_dir)
        self._current_project: Optional[Project] = None

        # Core services
        self._text_extractor = TextExtractor()
        self._llm_client = None
        self._report_worker = None
        self._analysis_worker = None
        self._extraction_worker = None
        self._index_worker = None

        # RAG Engine — lazy init nach erstem GUI-Render (verhindert Startup-Freeze
        # bei Remote-Ollama oder langsamer Netzwerkverbindung)
        self._rag_engine = None
        self._rag_engine_ready = False
        QTimer.singleShot(0, self._init_rag_engine)

        # Setup UI
        self._setup_window()
        self._setup_menu()
        self._setup_toolbar()
        self._setup_panels()
        self._setup_statusbar()

        # Connect signals
        self._connect_signals()

        # Create new project by default
        QTimer.singleShot(100, self._create_default_project)

    def _setup_window(self):
        """Configure the main window."""
        self.setWindowTitle("NoteSpaceLLM - Dokumenten-Analyse und Berichterstellung")
        self.setMinimumSize(1200, 800)

        # Try to restore geometry
        try:
            # Could load from settings file
            self.resize(1400, 900)
        except Exception as e:
            logger.warning("Could not restore window geometry: %s", e)

    def _setup_menu(self):
        """Set up the menu bar."""
        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu("&Datei")

        new_action = QAction("&Neues Projekt", self)
        new_action.setShortcut(QKeySequence.StandardKey.New)
        new_action.triggered.connect(self._new_project)
        file_menu.addAction(new_action)

        open_action = QAction("Projekt &Öffnen...", self)
        open_action.setShortcut(QKeySequence.StandardKey.Open)
        open_action.triggered.connect(self._open_project)
        file_menu.addAction(open_action)

        save_action = QAction("Projekt &speichern", self)
        save_action.setShortcut(QKeySequence.StandardKey.Save)
        save_action.triggered.connect(self._save_project)
        file_menu.addAction(save_action)

        file_menu.addSeparator()

        add_files_action = QAction("Dateien hinzufügen...", self)
        add_files_action.triggered.connect(self._add_files)
        file_menu.addAction(add_files_action)

        add_folder_action = QAction("Ordner hinzufügen...", self)
        add_folder_action.triggered.connect(self._add_folder)
        file_menu.addAction(add_folder_action)

        file_menu.addSeparator()

        export_action = QAction("&Exportieren...", self)
        export_action.setShortcut("Ctrl+E")
        export_action.triggered.connect(self._export)
        file_menu.addAction(export_action)

        workspace_export_action = QAction("Workspace exportieren (JSON)...", self)
        workspace_export_action.setShortcut("Ctrl+Shift+E")
        workspace_export_action.setToolTip(
            "Exportiert Projekt, Dokumente, Bericht und Chat als\n"
            "notespacellm-workspace-v1.json (datenschutzkonform)"
        )
        workspace_export_action.triggered.connect(self._export_workspace)
        file_menu.addAction(workspace_export_action)

        file_menu.addSeparator()

        exit_action = QAction("&Beenden", self)
        exit_action.setShortcut(QKeySequence.StandardKey.Quit)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Edit menu
        edit_menu = menubar.addMenu("&Bearbeiten")

        select_all_action = QAction("Alle auswählen", self)
        select_all_action.triggered.connect(self._select_all_docs)
        edit_menu.addAction(select_all_action)

        deselect_all_action = QAction("Alle abwählen", self)
        deselect_all_action.triggered.connect(self._deselect_all_docs)
        edit_menu.addAction(deselect_all_action)

        # LLM menu
        llm_menu = menubar.addMenu("&LLM")

        settings_action = QAction("LLM-&Einstellungen...", self)
        settings_action.setShortcut("Ctrl+L")
        settings_action.triggered.connect(self._llm_settings)
        llm_menu.addAction(settings_action)

        llm_menu.addSeparator()

        refresh_action = QAction("Modelle &aktualisieren", self)
        refresh_action.triggered.connect(self._refresh_models)
        llm_menu.addAction(refresh_action)

        # RAG menu
        rag_menu = menubar.addMenu("&RAG")

        index_all_action = QAction("Alle Dokumente indexieren", self)
        index_all_action.triggered.connect(self._index_all_documents)
        rag_menu.addAction(index_all_action)

        index_selected_action = QAction("Ausgewählte indexieren", self)
        index_selected_action.triggered.connect(self._index_selected_documents)
        rag_menu.addAction(index_selected_action)

        rag_menu.addSeparator()

        clear_index_action = QAction("Index leeren", self)
        clear_index_action.triggered.connect(self._clear_rag_index)
        rag_menu.addAction(clear_index_action)

        rag_menu.addSeparator()

        rag_stats_action = QAction("RAG-Statistiken...", self)
        rag_stats_action.triggered.connect(self._show_rag_stats)
        rag_menu.addAction(rag_stats_action)

        # Help menu
        help_menu = menubar.addMenu("&Hilfe")

        about_action = QAction("über NoteSpaceLLM", self)
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)

    def _setup_toolbar(self):
        """Set up the toolbar."""
        toolbar = QToolBar("Hauptwerkzeuge")
        toolbar.setMovable(False)
        self.addToolBar(toolbar)

        add_btn = QAction("+ Dateien", self)
        add_btn.setToolTip(
            "Dokumente hinzufügen\n"
            "Unterstützte Formate: PDF, DOCX, TXT, MD, XLSX, PPTX, ..."
        )
        add_btn.triggered.connect(self._add_files)
        toolbar.addAction(add_btn)

        folder_btn = QAction("+ Ordner", self)
        folder_btn.setToolTip("Ganzen Ordner rekursiv hinzufügen (alle unterstützten Dateitypen)")
        folder_btn.triggered.connect(self._add_folder)
        toolbar.addAction(folder_btn)

        toolbar.addSeparator()

        extract_btn = QAction("Text extrahieren", self)
        extract_btn.setToolTip(
            "Schritt 1 — Text aus Dokumenten extrahieren\n"
            "Läuft im Hintergrund, die GUI bleibt reaktionsfähig.\n"
            "Notwendig vor Analyse und Bericht-Erstellung."
        )
        extract_btn.triggered.connect(self._extract_all_text)
        toolbar.addAction(extract_btn)

        analyze_btn = QAction("Analysieren", self)
        analyze_btn.setToolTip(
            "Schritt 2 (optional) — Detailrecherchen per LLM\n"
            "Beantwortet die sub-queries pro Dokument (Rechtsklick auf Dokument).\n"
            "Kann lange dauern — läuft im Hintergrund."
        )
        analyze_btn.triggered.connect(self._run_analysis)
        toolbar.addAction(analyze_btn)

        toolbar.addSeparator()

        report_btn = QAction("Bericht erstellen", self)
        report_btn.setToolTip(
            "Schritt 3 — Bericht per LLM generieren\n"
            "Löst fehlende Textextraktion automatisch aus.\n"
            "Nutzt Hauptfragestellung + Detailanalysen aus dem Workflow-Panel."
        )
        report_btn.triggered.connect(self._generate_report)
        toolbar.addAction(report_btn)

        toolbar.addSeparator()

        help_btn = QAction("?", self)
        help_btn.setToolTip("Pipeline-Übersicht — Erklärt alle Schritte")
        help_btn.triggered.connect(self._show_pipeline_help)
        toolbar.addAction(help_btn)

    def _setup_panels(self):
        """Set up the main panels."""
        from .document_panel import DocumentPanel
        from .workflow_panel import WorkflowPanel
        from .chat_panel import ChatPanel
        from .output_panel import OutputPanel

        # Central widget with splitters
        central = QWidget()
        self.setCentralWidget(central)

        main_layout = QHBoxLayout(central)
        main_layout.setContentsMargins(5, 5, 5, 5)

        # Main horizontal splitter
        main_splitter = QSplitter(Qt.Orientation.Horizontal)

        # Left: Document panel
        self.document_panel = DocumentPanel()
        main_splitter.addWidget(self.document_panel)

        # Center: Workflow + Chat
        center_splitter = QSplitter(Qt.Orientation.Vertical)

        self.workflow_panel = WorkflowPanel()
        center_splitter.addWidget(self.workflow_panel)

        self.chat_panel = ChatPanel()
        center_splitter.addWidget(self.chat_panel)

        center_splitter.setSizes([400, 300])
        main_splitter.addWidget(center_splitter)

        # Right: Output panel
        self.output_panel = OutputPanel()
        main_splitter.addWidget(self.output_panel)

        # Set splitter sizes
        main_splitter.setSizes([300, 500, 400])

        main_layout.addWidget(main_splitter)

    def _setup_statusbar(self):
        """Set up the status bar."""
        self.statusbar = QStatusBar()
        self.setStatusBar(self.statusbar)
        self.statusbar.showMessage("Bereit")

        # Permanent pipeline phase indicator (right side of statusbar)
        self._pipeline_label = QLabel("📄 Keine Dokumente")
        self._pipeline_label.setStyleSheet("color: #888; font-size: 10px; padding: 0 8px;")
        self._pipeline_label.setToolTip(
            "Aktuelle Pipeline-Phase:\n"
            "📄 Keine Dokumente → Dateien hinzufügen\n"
            "🔤 Dokumente geladen → Text extrahieren (Schritt 1)\n"
            "✅ Extrahiert → Analysieren (opt.) oder Bericht erstellen (Schritt 3)\n"
            "🎯 Bereit → Bericht erstellen"
        )
        self.statusbar.addPermanentWidget(self._pipeline_label)

    def _connect_signals(self):
        """Connect panel signals."""
        # Document panel signals
        self.document_panel.selection_changed.connect(self._on_selection_changed)
        self.document_panel.subquery_requested.connect(self._on_subquery_requested)
        self.document_panel.files_added.connect(self._on_files_added)

        # Workflow panel signals
        self.workflow_panel.start_requested.connect(self._generate_report)

        # Chat panel signals
        self.chat_panel.message_sent.connect(self._on_chat_message)

        # Output panel signals
        self.output_panel.export_requested.connect(self._on_export_requested)
        self.output_panel.prompt_export_requested.connect(self._export_prompt)

    def _create_default_project(self):
        """Create a default project on startup."""
        self._current_project = self._project_manager.create_project(
            "Neues Projekt",
            "Was soll analysiert werden?",
            "analysis"
        )

        # Connect managers to panels
        self.document_panel.set_managers(
            self._current_project.documents,
            self._current_project.subqueries
        )

        # RAG-Engine mit Document Manager verbinden
        if self._rag_engine:
            self._current_project.documents.set_rag_engine(self._rag_engine)
            self.chat_panel.set_rag_engine(self._rag_engine)
            self.chat_panel.set_document_manager(self._current_project.documents)

        self.statusbar.showMessage("Neues Projekt erstellt")

    # Menu actions
    def _new_project(self):
        """Create a new project."""
        name, ok = QInputDialog.getText(self, "Neues Projekt", "Projektname:")
        if ok and name:
            self._project_manager.close_project()
            self._current_project = self._project_manager.create_project(name)

            self.document_panel.set_managers(
                self._current_project.documents,
                self._current_project.subqueries
            )

            # RAG-Engine verbinden
            if self._rag_engine:
                self._current_project.documents.set_rag_engine(self._rag_engine)
                self.chat_panel.set_rag_engine(self._rag_engine)
                self.chat_panel.set_document_manager(self._current_project.documents)

            self.statusbar.showMessage(f"Projekt '{name}' erstellt")

    def _open_project(self):
        """Open an existing project."""
        projects = self._project_manager.list_projects()
        if not projects:
            QMessageBox.information(self, "Projekt öffnen", "Keine Projekte vorhanden.")
            return

        names = [p["name"] for p in projects]
        name, ok = QInputDialog.getItem(
            self, "Projekt öffnen", "Projekt wählen:", names, editable=False
        )

        if ok and name:
            project = self._project_manager.open_project(name)
            if project:
                self._current_project = project
                self.document_panel.set_managers(
                    project.documents,
                    project.subqueries
                )
                self.workflow_panel.set_main_question(project.main_question)

                # RAG-Engine verbinden
                if self._rag_engine:
                    project.documents.set_rag_engine(self._rag_engine)
                    self.chat_panel.set_rag_engine(self._rag_engine)
                    self.chat_panel.set_document_manager(project.documents)

                self.statusbar.showMessage(f"Projekt '{name}' geöffnet")

    def _save_project(self):
        """Save the current project."""
        if self._project_manager.save_current():
            self.statusbar.showMessage("Projekt gespeichert")
        else:
            self.statusbar.showMessage("Fehler beim Speichern")

    def _add_files(self):
        """Add files to the project."""
        if not self._current_project:
            return

        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Dateien hinzufügen",
            "",
            "Alle unterstützten (*.pdf *.docx *.doc *.rtf *.txt *.md *.xlsx *.xls *.pptx *.py *.csv *.json *.xml *.eml *.msg);;Dokumente (*.pdf *.docx *.doc *.rtf *.txt *.md);;Tabellen (*.xlsx *.xls *.csv);;Code (*.py *.js *.java *.cpp *.c *.h);;Alle Dateien (*)"
        )

        if files:
            for f in files:
                self._current_project.documents.add_file(Path(f))
            self.statusbar.showMessage(f"{len(files)} Dateien hinzugefügt")
            self._update_pipeline_phase()

    def _add_folder(self):
        """Add a folder to the project."""
        if not self._current_project:
            return

        folder = QFileDialog.getExistingDirectory(self, "Ordner hinzufügen")
        if folder:
            docs = self._current_project.documents.add_directory(Path(folder))
            self.statusbar.showMessage(f"{len(docs)} Elemente hinzugefügt")

    def _export(self):
        """Export the report."""
        self._on_export_requested(
            self.output_panel.get_selected_formats(),
            str(self.output_panel.get_output_directory())
        )

    def _select_all_docs(self):
        """Select all documents."""
        if self._current_project:
            self._current_project.documents.select_all()

    def _deselect_all_docs(self):
        """Deselect all documents."""
        if self._current_project:
            self._current_project.documents.deselect_all()

    def _refresh_models(self):
        """Refresh the list of available models (async — no GUI freeze)."""
        if hasattr(self, '_model_load_worker') and self._model_load_worker and self._model_load_worker.isRunning():
            return

        self.statusbar.showMessage("Modelle werden abgefragt...")
        url = self._get_ollama_url()
        api_key = self._current_project.settings.ollama_api_key if self._current_project else ""

        self._model_load_worker = ModelLoadWorker(url, api_key)

        def _on_models(models, msg):
            self.statusbar.showMessage(f"Ollama: {msg}")
            self._model_load_worker = None

        self._model_load_worker.models_loaded.connect(_on_models)
        self._model_load_worker.start()

    def _llm_settings(self):
        """Show LLM settings dialog with provider, model, profiles and embedding config."""
        dialog = QDialog(self)
        dialog.setWindowTitle("LLM-Einstellungen")
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout(dialog)

        from ..core.app_config import get_app_config
        app_cfg = get_app_config()

        # ===== Profile Section =====
        profile_group = QGroupBox("Verbindungsprofile")
        profile_layout = QHBoxLayout(profile_group)

        profile_combo = QComboBox()
        profile_combo.setMinimumWidth(200)

        save_profile_btn = QPushButton("Speichern als...")
        rename_profile_btn = QPushButton("Umbenennen")
        delete_profile_btn = QPushButton("Löschen")

        profile_layout.addWidget(QLabel("Profil:"))
        profile_layout.addWidget(profile_combo, stretch=1)
        profile_layout.addWidget(save_profile_btn)
        profile_layout.addWidget(rename_profile_btn)
        profile_layout.addWidget(delete_profile_btn)

        layout.addWidget(profile_group)

        # ===== Provider Section =====
        provider_group = QGroupBox("Provider")
        provider_layout = QFormLayout(provider_group)

        provider_combo = QComboBox()
        provider_combo.addItems(["ollama", "openai", "anthropic", "claude-code"])
        provider_layout.addRow("Provider:", provider_combo)

        ollama_url_edit = QLineEdit()
        ollama_url_edit.setPlaceholderText("http://localhost:11434")
        provider_layout.addRow("Ollama URL:", ollama_url_edit)

        ollama_key_edit = QLineEdit()
        ollama_key_edit.setPlaceholderText("Leer = keine Auth (lokal)")
        ollama_key_edit.setEchoMode(QLineEdit.EchoMode.Password)
        provider_layout.addRow("API-Key:", ollama_key_edit)

        embedding_combo = QComboBox()
        embedding_combo.setEditable(True)
        embedding_combo.addItems(["nomic-embed-text", "mxbai-embed-large", "all-minilm"])
        provider_layout.addRow("Embedding-Modell:", embedding_combo)

        embedding_hint = QLabel("Muss auf dem Ollama-Server installiert sein (ollama pull ...)")
        embedding_hint.setStyleSheet("color: #888; font-size: 10px;")
        provider_layout.addRow("", embedding_hint)

        claude_mode_combo = QComboBox()
        claude_mode_combo.addItems(["api", "chat"])
        provider_layout.addRow("Claude Code Modus:", claude_mode_combo)

        claude_mode_hint = QLabel(
            "api = Hintergrund (Ergebnis zurück in App)\n"
            "chat = Konsole öffnet sich zum Weiter-Chatten"
        )
        claude_mode_hint.setStyleSheet("color: #888; font-size: 10px;")
        provider_layout.addRow("", claude_mode_hint)

        def _update_provider_visibility(provider: str):
            is_ollama = provider == "ollama"
            is_claude_code = provider == "claude-code"
            ollama_url_edit.setVisible(is_ollama)
            ollama_key_edit.setVisible(is_ollama)
            embedding_combo.setVisible(is_ollama)
            embedding_hint.setVisible(is_ollama)
            claude_mode_combo.setVisible(is_claude_code)
            claude_mode_hint.setVisible(is_claude_code)

        provider_combo.currentTextChanged.connect(_update_provider_visibility)

        layout.addWidget(provider_group)

        # ===== Model Section =====
        model_group = QGroupBox("Modell")
        model_layout = QVBoxLayout(model_group)

        model_combo = QComboBox()
        model_combo.setEditable(True)
        model_combo.setMinimumWidth(300)

        status_label = QLabel("")
        status_label.setStyleSheet("color: #666; font-size: 11px;")

        refresh_btn = QPushButton("Modelle laden")

        model_layout.addWidget(QLabel("Verfügbare Modelle:"))
        model_layout.addWidget(model_combo)

        btn_row = QHBoxLayout()
        btn_row.addWidget(refresh_btn)
        btn_row.addWidget(status_label)
        btn_row.addStretch()
        model_layout.addLayout(btn_row)

        layout.addWidget(model_group)

        api_hint = QLabel(
            "Für OpenAI/Anthropic: API-Key muss als Umgebungsvariable gesetzt sein\n"
            "(OPENAI_API_KEY bzw. ANTHROPIC_API_KEY)"
        )
        api_hint.setStyleSheet("color: #888; font-size: 10px;")
        layout.addWidget(api_hint)

        # Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        # ===== Helper: Populate form from values =====
        def _populate_form(provider, model, url, key, emb_model, cc_mode):
            provider_combo.setCurrentText(provider)
            ollama_url_edit.setText(url)
            ollama_key_edit.setText(key)
            embedding_combo.setCurrentText(emb_model)
            claude_mode_combo.setCurrentText(cc_mode)
            _update_provider_visibility(provider)

        # ===== Profile Logic =====
        def _refresh_profiles():
            profile_combo.clear()
            profile_combo.addItem("(Aktuelle Einstellungen)")
            for name in app_cfg.list_profile_names():
                profile_combo.addItem(name)
            # Select active profile
            active = app_cfg.active_profile
            if active:
                idx = profile_combo.findText(active)
                if idx >= 0:
                    profile_combo.setCurrentIndex(idx)

        def _on_profile_selected(name: str):
            if name == "(Aktuelle Einstellungen)" or not name:
                return
            profiles = app_cfg.profiles
            if name in profiles:
                p = profiles[name]
                _populate_form(
                    p.get("llm_provider", "ollama"),
                    p.get("llm_model", "llama3"),
                    p.get("ollama_base_url", "http://localhost:11434"),
                    p.get("ollama_api_key", ""),
                    p.get("embedding_model", "nomic-embed-text"),
                    app_cfg.claude_code_mode,
                )
                # Trigger model reload
                load_models_for_provider(p.get("llm_provider", "ollama"))
                # Try select model after load
                idx = model_combo.findText(p.get("llm_model", ""))
                if idx >= 0:
                    model_combo.setCurrentIndex(idx)
                elif model_combo.count() > 0:
                    model_combo.setCurrentText(p.get("llm_model", ""))

        def _save_profile_dialog():
            name, ok = QInputDialog.getText(dialog, "Profil speichern", "Profilname:")
            if ok and name:
                # Save current form values as profile
                app_cfg.llm_provider = provider_combo.currentText()
                app_cfg.llm_model = model_combo.currentText().strip()
                app_cfg.ollama_base_url = ollama_url_edit.text().strip() or "http://localhost:11434"
                app_cfg.ollama_api_key = ollama_key_edit.text().strip()
                app_cfg.embedding_model = embedding_combo.currentText().strip() or "nomic-embed-text"
                app_cfg.save_profile(name)
                _refresh_profiles()
                profile_combo.setCurrentText(name)
                status_label.setText(f"Profil '{name}' gespeichert")

        def _delete_profile():
            name = profile_combo.currentText()
            if name == "(Aktuelle Einstellungen)":
                return
            if app_cfg.delete_profile(name):
                _refresh_profiles()
                status_label.setText(f"Profil '{name}' gelöscht")

        def _rename_profile():
            old_name = profile_combo.currentText()
            if old_name == "(Aktuelle Einstellungen)":
                return
            new_name, ok = QInputDialog.getText(
                dialog, "Profil umbenennen",
                "Neuer Profilname:",
                text=old_name
            )
            if ok and new_name.strip() and new_name.strip() != old_name:
                if app_cfg.rename_profile(old_name, new_name.strip()):
                    _refresh_profiles()
                    profile_combo.setCurrentText(new_name.strip())
                    status_label.setText(f"'{old_name}' → '{new_name.strip()}'")
                else:
                    status_label.setText(f"Umbenennen fehlgeschlagen (Name belegt?)")

        profile_combo.currentTextChanged.connect(_on_profile_selected)
        save_profile_btn.clicked.connect(_save_profile_dialog)
        rename_profile_btn.clicked.connect(_rename_profile)
        delete_profile_btn.clicked.connect(_delete_profile)

        # ===== Model Loading =====
        OPENAI_MODELS = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"]
        ANTHROPIC_MODELS = [
            "claude-sonnet-4-6", "claude-haiku-4-5-20251001",
            "claude-3-5-sonnet-20241022", "claude-3-haiku-20240307"
        ]

        _settings_model_worker = [None]  # mutable container for worker ref

        def _select_current_model():
            cur = app_cfg.llm_model
            idx = model_combo.findText(cur)
            if idx >= 0:
                model_combo.setCurrentIndex(idx)
            elif model_combo.count() > 0:
                model_combo.setCurrentIndex(0)

        def load_models_for_provider(provider: str):
            model_combo.clear()
            status_label.setText("Lade...")

            if provider == "ollama":
                # Async model load — no GUI freeze
                if _settings_model_worker[0] and _settings_model_worker[0].isRunning():
                    return
                url = ollama_url_edit.text().strip() or "http://localhost:11434"
                key = ollama_key_edit.text().strip()
                worker = ModelLoadWorker(url, key)

                def _on_ollama_models(models, msg):
                    model_combo.clear()
                    if models:
                        model_combo.addItems(models)
                    status_label.setText(msg)
                    _select_current_model()
                    _settings_model_worker[0] = None

                worker.models_loaded.connect(_on_ollama_models)
                _settings_model_worker[0] = worker
                worker.start()
                return  # combo will be filled once worker finishes

            elif provider == "openai":
                model_combo.addItems(OPENAI_MODELS)
                status_label.setText("Standard-Modelle (editierbar)")

            elif provider == "anthropic":
                model_combo.addItems(ANTHROPIC_MODELS)
                status_label.setText("Standard-Modelle (editierbar)")

            elif provider == "claude-code":
                import shutil
                claude_available = shutil.which("claude") is not None
                model_combo.addItems(["sonnet", "opus", "haiku"])
                if claude_available:
                    status_label.setText("Claude Code CLI gefunden")
                else:
                    status_label.setText("Claude Code CLI NICHT gefunden!")

            _select_current_model()

        provider_combo.currentTextChanged.connect(load_models_for_provider)
        refresh_btn.clicked.connect(lambda: load_models_for_provider(provider_combo.currentText()))

        # ===== Initial populate =====
        _populate_form(
            app_cfg.llm_provider, app_cfg.llm_model,
            app_cfg.ollama_base_url, app_cfg.ollama_api_key,
            app_cfg.embedding_model, app_cfg.claude_code_mode,
        )
        _refresh_profiles()
        load_models_for_provider(app_cfg.llm_provider)

        # ===== Dialog Accept =====
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_provider = provider_combo.currentText()
            new_model = model_combo.currentText().strip()

            if not new_model:
                QMessageBox.warning(self, "LLM", "Kein Modell ausgewählt.")
                return

            new_url = ollama_url_edit.text().strip() or "http://localhost:11434"
            new_key = ollama_key_edit.text().strip()
            new_embedding = embedding_combo.currentText().strip() or "nomic-embed-text"
            new_claude_mode = claude_mode_combo.currentText()

            if self._current_project:
                self._current_project.settings.llm_provider = new_provider
                self._current_project.settings.llm_model = new_model
                self._current_project.settings.ollama_base_url = new_url
                self._current_project.settings.ollama_api_key = new_key

            app_cfg.llm_provider = new_provider
            app_cfg.llm_model = new_model
            app_cfg.ollama_base_url = new_url
            app_cfg.ollama_api_key = new_key
            app_cfg.embedding_model = new_embedding
            app_cfg.claude_code_mode = new_claude_mode
            app_cfg.save()

            # Re-init RAG engine with new embedding model
            self._init_rag_engine()
            if self._rag_engine and self._current_project:
                self._current_project.documents.set_rag_engine(self._rag_engine)
                self.chat_panel.set_rag_engine(self._rag_engine)
                # Nachindexierung: Dokumente mit Text aber ohne Index
                unindexed = [d for d in self._current_project.documents.documents
                             if d.extracted_text and not d.is_indexed and not d.is_directory]
                if unindexed:
                    self._start_index_worker([(d.id, d.name) for d in unindexed])

            self._init_llm_client()
            self.statusbar.showMessage(f"LLM: {new_provider} / {new_model}")

    def _show_about(self):
        """Show about dialog."""
        QMessageBox.about(
            self,
            "über NoteSpaceLLM",
            "NoteSpaceLLM v1.0.0\n\n"
            "Ein privater NotebookLM-Clone für lokale Dokumentenanalyse "
            "und Berichterstellung.\n\n"
            "Features:\n"
            "- Drag & Drop Dokumentenverwaltung\n"
            "- Detailrecherchen pro Dokument\n"
            "- Visuelle Workflow-Steuerung\n"
            "- LLM-gestützter Dokumenten-Chat\n"
            "- Multi-Format-Export"
        )

    # Panel callbacks
    def _on_selection_changed(self):
        """Handle document selection change."""
        self._update_document_context()

    def _on_files_added(self):
        """Handle files added — start async extraction for pending docs."""
        if not self._current_project:
            return
        pending = self._current_project.documents.pop_pending_extractions()
        if not pending:
            return

        # Skip if extraction already running — queue will be picked up next time
        if self._extraction_worker and self._extraction_worker.isRunning():
            return

        self.statusbar.showMessage(f"Extrahiere Text aus {len(pending)} Dokumenten...")

        from ..core.text_extractor import TextExtractor
        if not hasattr(self, '_text_extractor') or self._text_extractor is None:
            self._text_extractor = TextExtractor()

        self._extraction_worker = ExtractionWorker(self._text_extractor, pending)

        def _on_doc_extracted(doc_id, text, error):
            from ..core.document_manager import DocumentStatus
            if error:
                self._current_project.documents.set_status(doc_id, DocumentStatus.ERROR, error)
            else:
                self._current_project.documents.update_content(doc_id, text)

        def _on_extraction_complete():
            self._extraction_worker = None
            self._update_document_context()
            self._update_pipeline_phase()
            self.statusbar.showMessage(f"Extraktion abgeschlossen: {len(pending)} Dokumente")

        self._extraction_worker.doc_extracted.connect(_on_doc_extracted)
        self._extraction_worker.progress.connect(
            lambda cur, total, name: self.statusbar.showMessage(
                f"Extrahiere ({cur+1}/{total}): {name}") if name else None)
        self._extraction_worker.all_complete.connect(_on_extraction_complete)
        self._extraction_worker.start()

    def _on_subquery_requested(self, doc_id: str, query_type: str, query_text: str):
        """Handle sub-query request."""
        self.statusbar.showMessage(f"Detailrecherche hinzugefügt: {query_text[:50]}...")

    def _on_chat_message(self, message: str):
        """Handle chat message."""
        # Context is already set via _update_document_context
        pass

    def _on_export_requested(self, formats: list, directory: str):
        """Handle export request — lets user choose filename via dialog."""
        content = self.output_panel.get_content()
        if not content:
            QMessageBox.warning(self, "Export", "Kein Inhalt zum Exportieren.")
            return

        project_name = self._current_project.name if self._current_project else "report"
        safe_name = "".join(c for c in project_name if c.isalnum() or c in " -_").strip().replace(" ", "_")

        # Build list of (format, path) pairs — ask user for filename
        _filter_map = {
            "md": "Markdown (*.md)", "txt": "Text (*.txt)",
            "html": "HTML (*.html)", "pdf": "PDF (*.pdf)", "docx": "Word (*.docx)"
        }
        if len(formats) == 1:
            fmt = formats[0]
            default_path = Path(directory) / f"{safe_name}.{fmt}"
            filepath, _ = QFileDialog.getSaveFileName(
                self, "Bericht exportieren",
                str(default_path),
                f"{_filter_map.get(fmt, 'Dateien (*)')};;Alle Dateien (*)"
            )
            if not filepath:
                return
            formats_and_paths = [(fmt, Path(filepath))]
        else:
            base_name, ok = QInputDialog.getText(
                self, "Exportdatei-Name",
                "Basis-Dateiname (ohne Erweiterung):",
                text=safe_name
            )
            if not ok or not base_name.strip():
                return
            output_dir = Path(directory)
            output_dir.mkdir(parents=True, exist_ok=True)
            formats_and_paths = [(fmt, output_dir / f"{base_name.strip()}.{fmt}") for fmt in formats]

        exported = []
        for fmt, filepath in formats_and_paths:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            try:
                if fmt == "md":
                    filepath.write_text(content, encoding="utf-8")
                    exported.append(filepath.name)

                elif fmt == "txt":
                    import re
                    plain = re.sub(r'[#*`_]', '', content)
                    filepath.write_text(plain, encoding="utf-8")
                    exported.append(filepath.name)

                elif fmt == "html":
                    html = self._md_to_html(content)
                    filepath.write_text(html, encoding="utf-8")
                    exported.append(filepath.name)

                elif fmt == "pdf":
                    if self._export_pdf(content, filepath):
                        exported.append(filepath.name)

                elif fmt == "docx":
                    if self._export_docx(content, filepath):
                        exported.append(filepath.name)

            except Exception as e:
                self.statusbar.showMessage(f"Fehler bei {fmt}: {e}")

        if exported:
            out_dir = str(formats_and_paths[0][1].parent)
            self.output_panel.set_status(f"Exportiert: {', '.join(exported)}")
            QMessageBox.information(
                self, "Export",
                f"Erfolgreich exportiert:\n{chr(10).join(exported)}\n\nOrdner: {out_dir}"
            )

    def _md_to_html(self, content: str) -> str:
        """Convert markdown to simple HTML."""
        import re

        html = content
        # Headers
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        # Bold
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        # Italic
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        # Code
        html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)
        # Paragraphs
        html = re.sub(r'\n\n', r'</p><p>', html)

        return f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Report</title></head><body><p>{html}</p></body></html>"

    def _export_pdf(self, content: str, filepath: Path) -> bool:
        """Export to PDF."""
        try:
            # Try markdown2pdf or pandoc
            import subprocess
            import tempfile

            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
                f.write(content)
                md_path = f.name

            result = subprocess.run(
                ['pandoc', md_path, '-o', str(filepath)],
                capture_output=True,
                timeout=60
            )

            Path(md_path).unlink()
            return result.returncode == 0

        except Exception:
            return False

    def _export_docx(self, content: str, filepath: Path) -> bool:
        """Export to DOCX."""
        try:
            from docx import Document

            doc = Document()
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)

            doc.save(str(filepath))
            return True

        except ImportError:
            return False

    # Core functionality
    def _get_ollama_url(self) -> str:
        """Get the Ollama base URL from current project or app config."""
        if self._current_project:
            return self._current_project.settings.ollama_base_url
        from ..core.app_config import get_app_config
        return get_app_config().ollama_base_url

    def _init_llm_client(self):
        """Initialize the LLM client."""
        if not self._current_project:
            return

        from ..llm import create_llm_client

        provider = self._current_project.settings.llm_provider
        model = self._current_project.settings.llm_model
        base_url = self._current_project.settings.ollama_base_url
        api_key = self._current_project.settings.ollama_api_key

        try:
            from ..core.app_config import get_app_config
            claude_mode = get_app_config().claude_code_mode
            self._llm_client = create_llm_client(
                provider, model, base_url=base_url,
                api_key=api_key, claude_code_mode=claude_mode
            )
            self.chat_panel.set_llm_client(self._llm_client)
            self.statusbar.showMessage(f"LLM initialisiert: {provider}/{model}")
        except Exception as e:
            self.statusbar.showMessage(f"LLM-Fehler: {e}")

    def _update_pipeline_phase(self):
        """Update the pipeline phase indicator in the status bar."""
        if not hasattr(self, '_pipeline_label'):
            return
        if not self._current_project:
            self._pipeline_label.setText("📄 Keine Dokumente")
            return

        docs = [d for d in self._current_project.documents.documents if not d.is_directory]
        if not docs:
            self._pipeline_label.setText("📄 Keine Dokumente")
            return

        extracted = [d for d in docs if d.extracted_text]
        indexed = [d for d in docs if d.is_indexed]

        if len(extracted) == 0:
            self._pipeline_label.setText(f"🔤 {len(docs)} Dok. → Text extrahieren (Schritt 1)")
        elif len(extracted) < len(docs):
            self._pipeline_label.setText(
                f"⏳ {len(extracted)}/{len(docs)} extrahiert → weiter extrahieren"
            )
        elif indexed:
            self._pipeline_label.setText(
                f"🎯 {len(extracted)} extrahiert, {len(indexed)} indexiert → Bericht erstellen"
            )
        else:
            self._pipeline_label.setText(
                f"✅ {len(extracted)} extrahiert → Analysieren oder Bericht erstellen"
            )

    def _update_document_context(self):
        """Update the document context for chat."""
        if not self._current_project:
            return

        docs = self._current_project.documents.selected_documents
        context_parts = []

        for doc in docs:
            if doc.extracted_text:
                context_parts.append(f"--- {doc.name} ---\n{doc.extracted_text}")

        context = "\n\n".join(context_parts)
        self.chat_panel.set_document_context(context)
        self._update_pipeline_phase()

    def _extract_all_text(self, on_complete=None):
        """Extract text from all documents in background thread."""
        if not self._current_project:
            if on_complete:
                on_complete()
            return

        docs = self._current_project.documents.selected_documents
        if not docs:
            if on_complete:
                on_complete()
            return

        # Nur Dokumente ohne extrahierten Text
        from ..core.document_manager import DocumentStatus
        docs_to_extract = [d for d in docs if not d.extracted_text and not d.is_directory]

        if not docs_to_extract:
            if on_complete:
                on_complete()
            return

        tasks = [(d.id, d.path, d.name) for d in docs_to_extract]
        total = len(tasks)
        self.statusbar.showMessage(f"Extrahiere Text aus {total} Dokumenten...")

        for d in docs_to_extract:
            self._current_project.documents.set_status(d.id, DocumentStatus.EXTRACTING)

        self._extraction_worker = ExtractionWorker(self._text_extractor, tasks)

        def _on_doc_extracted(doc_id, text, error):
            if error:
                self._current_project.documents.set_status(doc_id, DocumentStatus.ERROR, error)
            else:
                self._current_project.documents.update_content(doc_id, text)

        def _on_extraction_progress(current, total_count, filename):
            if filename:
                self.statusbar.showMessage(f"Extrahiere ({current+1}/{total_count}): {filename}")

        def _on_extraction_complete():
            self._update_document_context()
            self.statusbar.showMessage(f"Textextraktion abgeschlossen: {total} Dokumente")
            self._extraction_worker = None
            if on_complete:
                on_complete()

        self._extraction_worker.doc_extracted.connect(_on_doc_extracted)
        self._extraction_worker.progress.connect(_on_extraction_progress)
        self._extraction_worker.all_complete.connect(_on_extraction_complete)
        self._extraction_worker.start()

    def _run_analysis(self):
        """Run sub-query analyses in background thread."""
        # Prüfen ob bereits eine Analyse läuft
        if self._analysis_worker and self._analysis_worker.isRunning():
            self.statusbar.showMessage("Analyse läuft bereits...")
            return

        if not self._current_project or not self._llm_client:
            self._init_llm_client()
            if not self._llm_client:
                QMessageBox.warning(self, "Analyse", "Kein LLM-Client verfügbar.")
                return

        queries = self._current_project.subqueries.pending_queries
        if not queries:
            self.statusbar.showMessage("Keine ausstehenden Analysen")
            return

        # Build task list: [(query_id, prompt)] for the worker
        tasks = []
        for query in queries:
            doc = self._current_project.documents.get_document(query.document_id)
            if not doc or not doc.extracted_text:
                self._current_project.subqueries.set_error(query.id, "Dokument nicht verfügbar")
                continue
            self._current_project.subqueries.set_running(query.id)
            prompt = query.build_prompt(doc.extracted_text)
            tasks.append((query.id, prompt))

        if not tasks:
            return

        total = len(tasks)
        self.statusbar.showMessage(f"Führe {total} Analysen durch...")

        self._analysis_worker = AnalysisWorker(self._llm_client, tasks)
        self._analysis_worker.query_complete.connect(self._on_analysis_result)
        self._analysis_worker.all_complete.connect(
            lambda: self.statusbar.showMessage(f"Analysen abgeschlossen: {total}"))
        self._analysis_worker.start()

    def _generate_report(self):
        """Generate the main report (fully async, no GUI freeze)."""
        if not self._current_project:
            return

        # Prüfen ob bereits ein Bericht generiert wird
        if hasattr(self, '_report_worker') and self._report_worker and self._report_worker.isRunning():
            QMessageBox.information(
                self, "Bericht",
                "Ein Bericht wird bereits erstellt. Bitte warten."
            )
            return

        # Initialize LLM if needed
        if not self._llm_client:
            self._init_llm_client()
            if not self._llm_client:
                QMessageBox.warning(self, "Bericht", "Kein LLM-Client verfügbar.")
                return

        self.output_panel.clear_content()
        self.output_panel.set_status("Vorbereitung: Extrahiere Texte...")
        self.statusbar.showMessage("Bericht wird vorbereitet...")

        # Step 1: Extract text (async) -> Step 2: Build & send prompt
        self._extract_all_text(on_complete=self._generate_report_step2)

    def _generate_report_step2(self):
        """Build prompt and start LLM generation (called after extraction)."""
        # Build the main prompt
        main_question = self.workflow_panel.get_main_question()
        if not main_question:
            main_question = "Erstelle einen umfassenden Analysebericht."

        # Collect document content
        docs = self._current_project.documents.selected_documents
        doc_content = "\n\n".join([
            f"=== {doc.name} ===\n{doc.extracted_text}"
            for doc in docs if doc.extracted_text
        ])

        if not doc_content.strip():
            self.output_panel.set_status("Keine Dokumenteninhalte verfügbar.")
            QMessageBox.warning(self, "Bericht", "Keine extrahierten Texte vorhanden.")
            return

        # Collect sub-query results
        subquery_results = self._current_project.subqueries.get_results_for_report()
        subquery_content = ""
        for doc_id, results in subquery_results.items():
            doc = self._current_project.documents.get_document(doc_id)
            doc_name = doc.name if doc else doc_id
            for r in results:
                subquery_content += f"\n--- Detailanalyse: {doc_name} ({r['type']}) ---\n"
                subquery_content += f"Frage: {r['query']}\n"
                subquery_content += f"Ergebnis: {r['result']}\n"

        # Build final prompt
        prompt = f"""Du bist ein erfahrener Analyst und erstellst professionelle Berichte.

HAUPTFRAGESTELLUNG:
{main_question}

DOKUMENTENINHALTE:
{doc_content[:80000]}

DETAILANALYSEN:
{subquery_content[:20000]}

AUFGABE:
Erstelle einen strukturierten, professionellen Bericht basierend auf den Dokumenten.
Der Bericht soll die Hauptfragestellung beantworten und die Erkenntnisse aus den
Detailanalysen integrieren.

Strukturiere den Bericht mit:
1. Zusammenfassung (Executive Summary)
2. Einleitung und Fragestellung
3. Analyse der Kernthemen
4. Detailergebnisse
5. Schlussfolgerungen und Empfehlungen

Verwende Markdown-Formatierung."""

        # Update workflow status
        workflow = self.workflow_panel.get_current_workflow()
        if workflow:
            for step in workflow.steps:
                self.workflow_panel.update_step_status(step.id, "running")

        # Generate report in background thread
        self.output_panel.set_status("Generiere Bericht...")

        from .chat_panel import LLMWorker
        self._report_worker = LLMWorker(self._llm_client, prompt, "")
        self._report_workflow = workflow

        self._report_worker.response_chunk.connect(
            lambda chunk: self.output_panel.append_content(chunk))
        self._report_worker.response_complete.connect(self._on_report_complete)
        self._report_worker.error_occurred.connect(self._on_report_error)
        self._report_worker.start()

    def _on_analysis_result(self, query_id: str, response: str, error: str):
        """Handle a single analysis result from the worker."""
        if error:
            self._current_project.subqueries.set_error(query_id, error)
        else:
            self._current_project.subqueries.set_result(query_id, response)

    def _on_report_complete(self, full_response: str):
        """Handle completed report generation."""
        self.output_panel.set_status("Bericht erstellt")
        if self._report_workflow:
            for step in self._report_workflow.steps:
                self.workflow_panel.update_step_status(step.id, "completed")
        self.statusbar.showMessage("Bericht erfolgreich erstellt")
        self._report_worker = None

    def _on_report_error(self, error: str):
        """Handle report generation error."""
        self.output_panel.set_status(f"Fehler: {error}")
        self.statusbar.showMessage(f"Fehler bei Berichterstellung: {error}")
        self._report_worker = None

    def _export_prompt(self):
        """Export the analysis prompt as .md file for manual LLM usage."""
        if not self._current_project:
            return

        # Collect document content
        docs = self._current_project.documents.selected_documents
        doc_content = "\n\n".join([
            f"=== {doc.name} ===\n{doc.extracted_text}"
            for doc in docs if doc.extracted_text
        ])

        if not doc_content:
            QMessageBox.warning(self, "Prompt-Export",
                                "Keine Dokumente mit extrahiertem Text vorhanden.\n"
                                "Bitte zuerst Dokumente hinzufügen.")
            return

        main_question = self.workflow_panel.get_main_question()
        if not main_question:
            main_question = "Erstelle einen umfassenden Analysebericht."

        # Collect sub-query results
        subquery_results = self._current_project.subqueries.get_results_for_report()
        subquery_content = ""
        for doc_id, results in subquery_results.items():
            doc = self._current_project.documents.get_document(doc_id)
            doc_name = doc.name if doc else doc_id
            for r in results:
                subquery_content += f"\n--- Detailanalyse: {doc_name} ({r['type']}) ---\n"
                subquery_content += f"Frage: {r['query']}\nErgebnis: {r['result']}\n"

        # Build prompt
        from ..llm.claude_code_client import ClaudeCodeClient
        from datetime import datetime

        prompt = f"""Du bist ein erfahrener Analyst und erstellst professionelle Berichte.

HAUPTFRAGESTELLUNG:
{main_question}

DOKUMENTENINHALTE:
{doc_content[:80000]}

DETAILANALYSEN:
{subquery_content[:20000] if subquery_content else '(keine)'}

AUFGABE:
Erstelle einen strukturierten, professionellen Bericht basierend auf den Dokumenten.
Der Bericht soll die Hauptfragestellung beantworten und die Erkenntnisse aus den
Detailanalysen integrieren.

Strukturiere den Bericht mit:
1. Zusammenfassung (Executive Summary)
2. Einleitung und Fragestellung
3. Analyse der Kernthemen
4. Detailergebnisse
5. Schlussfolgerungen und Empfehlungen

Verwende Markdown-Formatierung."""

        # Ask where to save
        project_name = self._current_project.name
        safe_name = "".join(c for c in project_name if c.isalnum() or c in " -_").strip().replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"prompt_{safe_name}_{timestamp}.md"

        filepath, _ = QFileDialog.getSaveFileName(
            self, "Prompt exportieren",
            str(self.output_panel.get_output_directory() / default_name),
            "Markdown (*.md);;Text (*.txt);;Alle Dateien (*)"
        )

        if filepath:
            filepath = Path(filepath)
            ClaudeCodeClient.export_prompt(prompt, "", filepath)
            self.output_panel.set_status(f"Prompt exportiert: {filepath.name}")
            self.statusbar.showMessage(f"Prompt exportiert: {filepath}")

    def _show_pipeline_help(self):
        """Öffnet den Pipeline-Hilfe-Dialog."""
        dialog = _PipelineHelpDialog(self)
        dialog.exec()

    def _export_workspace(self):
        """Exportiert das aktuelle Projekt als notespacellm-workspace-v1.json."""
        if not self._current_project:
            QMessageBox.warning(self, "Workspace exportieren", "Kein Projekt geöffnet.")
            return

        from ..core.workspace_exporter import build_workspace_export_payload, export_workspace_to_file
        from datetime import datetime

        project_name = self._current_project.name
        safe_name = "".join(c for c in project_name if c.isalnum() or c in " -_").strip().replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"{safe_name}_workspace_{timestamp}.json"

        filepath, _ = QFileDialog.getSaveFileName(
            self,
            "Workspace exportieren",
            str(self.output_panel.get_output_directory() / default_name),
            "JSON (*.json);;Alle Dateien (*)"
        )
        if not filepath:
            return

        try:
            report_content = self.output_panel.get_content()
            raw_msgs = self.chat_panel.get_messages() if hasattr(self.chat_panel, 'get_messages') else []
            chat_history = [{"role": m.role, "content": m.content} for m in raw_msgs]

            payload = build_workspace_export_payload(
                project=self._current_project,
                report_content=report_content,
                chat_history=chat_history,
                include_subquery_excerpts=True,
            )
            export_workspace_to_file(payload, Path(filepath))

            doc_count = len([d for d in self._current_project.documents.documents
                             if not d.is_directory])
            self.statusbar.showMessage(
                f"Workspace exportiert: {Path(filepath).name} ({doc_count} Dokumente)"
            )
            QMessageBox.information(
                self, "Workspace exportiert",
                f"Erfolgreich exportiert:\n{Path(filepath).name}\n\n"
                f"Enthält: {doc_count} Dokument(e), Bericht, Chat\n"
                f"API-Schlüssel und Rohdaten wurden nicht exportiert."
            )
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Export fehlgeschlagen:\n{e}")
            self.statusbar.showMessage(f"Workspace-Export fehlgeschlagen: {e}")

    def closeEvent(self, event):
        """Handle window close."""
        # Save project
        self._project_manager.close_project()
        event.accept()

    # ==================== RAG Methods ====================

    def _init_rag_engine(self):
        """Initialize the RAG engine with ChromaDB and Ollama Embeddings.

        Wird via QTimer.singleShot(0) aufgerufen — erst nach dem ersten GUI-Render,
        damit der Start nicht durch Netzwerk-Latenz oder schwere Imports blockiert wird.
        """
        if hasattr(self, "statusbar"):
            self.statusbar.showMessage("RAG-Engine wird verbunden …")

        try:
            from ..rag.engine import RAGEngine
            from ..core.app_config import get_app_config

            app_cfg = get_app_config()
            ollama_url = app_cfg.ollama_base_url
            llm_model = app_cfg.llm_model or "llama3.2"
            embedding_model = app_cfg.embedding_model or "nomic-embed-text"

            # RAG nur initialisieren wenn Provider ollama ist
            # (Embeddings brauchen Ollama)
            if app_cfg.llm_provider not in ("ollama",):
                logger.info("RAG Engine übersprungen (Provider ist nicht Ollama)")
                self._rag_engine = None
                if hasattr(self, "statusbar"):
                    self.statusbar.showMessage("Bereit (RAG nicht aktiv — kein Ollama-Provider)")
                return

            # Storage-Verzeichnis
            storage_dir = Path.home() / "NoteSpaceLLM" / "storage" / "chroma_db"
            storage_dir.parent.mkdir(parents=True, exist_ok=True)

            self._rag_engine = RAGEngine(
                persist_directory=str(storage_dir),
                collection_name="notespace_documents",
                embedding_model=embedding_model,
                llm_model=llm_model,
                ollama_base_url=ollama_url,
                api_key=app_cfg.ollama_api_key
            )
            self._rag_engine_ready = True
            logger.info(f"RAG Engine initialisiert (URL: {ollama_url}, Embedding: {embedding_model})")
            if hasattr(self, "statusbar"):
                self.statusbar.showMessage(f"RAG-Engine bereit ({embedding_model} @ {ollama_url})", 5000)

        except Exception as e:
            logger.error(f"RAG Engine Initialisierung fehlgeschlagen: {e}")
            self._rag_engine = None
            if hasattr(self, "statusbar"):
                self.statusbar.showMessage(f"RAG-Engine nicht verfügbar: {e}", 8000)

    def _index_all_documents(self):
        """Index all documents with extracted text (async)."""
        if not self._current_project or not self._rag_engine:
            QMessageBox.warning(self, "RAG", "RAG Engine nicht verfügbar.")
            return

        if hasattr(self, '_index_worker') and self._index_worker and self._index_worker.isRunning():
            QMessageBox.information(self, "RAG", "Indexierung läuft bereits.")
            return

        docs = self._current_project.documents.documents
        docs_with_text = [d for d in docs if d.extracted_text and not d.is_directory]

        if not docs_with_text:
            QMessageBox.information(self, "RAG", "Keine Dokumente mit extrahiertem Text gefunden.")
            return

        self._start_index_worker([(d.id, d.name) for d in docs_with_text])

    def _index_selected_documents(self):
        """Index only selected documents (async)."""
        if not self._current_project or not self._rag_engine:
            QMessageBox.warning(self, "RAG", "RAG Engine nicht verfügbar.")
            return

        if hasattr(self, '_index_worker') and self._index_worker and self._index_worker.isRunning():
            QMessageBox.information(self, "RAG", "Indexierung läuft bereits.")
            return

        docs = self._current_project.documents.selected_documents
        docs_with_text = [d for d in docs if d.extracted_text and not d.is_directory]

        if not docs_with_text:
            QMessageBox.information(self, "RAG", "Keine ausgewählten Dokumente mit Text gefunden.")
            return

        self._start_index_worker([(d.id, d.name) for d in docs_with_text])

    def _start_index_worker(self, doc_ids_and_names):
        """Start the index worker thread."""
        self.statusbar.showMessage(f"Indexiere {len(doc_ids_and_names)} Dokumente...")

        self._index_worker = IndexWorker(self._current_project.documents, doc_ids_and_names)

        self._index_worker.progress.connect(
            lambda cur, total, name: self.statusbar.showMessage(
                f"Indexiere ({cur+1}/{total}): {name}") if name else None)
        self._index_worker.all_complete.connect(self._on_index_complete)
        self._index_worker.start()

    def _on_index_complete(self, indexed, total):
        """Handle index worker completion."""
        self._index_worker = None
        self.statusbar.showMessage(f"RAG: {indexed}/{total} Dokumente indexiert")
        QMessageBox.information(self, "RAG", f"{indexed}/{total} Dokumente erfolgreich indexiert.")

    def _clear_rag_index(self):
        """Clear the RAG index."""
        if not self._rag_engine:
            return

        reply = QMessageBox.question(
            self, "Index leeren",
            "Möchtest du den gesamten RAG-Index leeren?\n"
            "Alle Embeddings werden gelöscht.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            if self._rag_engine.clear_index():
                # Reset indexed status
                if self._current_project:
                    for doc in self._current_project.documents.documents:
                        doc.is_indexed = False
                        doc.chunk_count = 0

                self.statusbar.showMessage("RAG-Index geleert")
                QMessageBox.information(self, "RAG", "Index erfolgreich geleert.")
            else:
                QMessageBox.warning(self, "RAG", "Fehler beim Leeren des Index.")

    def _show_rag_stats(self):
        """Show RAG statistics dialog."""
        if not self._rag_engine:
            QMessageBox.warning(self, "RAG", "RAG Engine nicht verfügbar.")
            return

        stats = self._rag_engine.get_statistics()
        connection = self._rag_engine.test_connection()

        # Dokumenten-Statistiken
        doc_stats = {}
        if self._current_project:
            doc_stats = self._current_project.documents.get_rag_statistics()

        info = f"""RAG Engine Statistiken
========================

Embedding-Modell: {stats.get('embedding_model', 'N/A')}
LLM-Modell: {stats.get('llm_model', 'N/A')}
Collection: {stats.get('collection_name', 'N/A')}
Persist-Verzeichnis: {stats.get('persist_directory', 'N/A')}

Indexierte Chunks: {stats.get('total_chunks', 0)}
Indexierte Dokumente: {doc_stats.get('indexed_documents', 0)}

Verbindungsstatus:
- Embeddings: {'✅' if connection.get('embeddings') else '❌'}
- Vector Store: {'✅' if connection.get('vectorstore') else '❌'}
- LLM: {'✅' if connection.get('llm') else '❌'}

Auto-Indexierung: {'Aktiv' if doc_stats.get('auto_index', False) else 'Inaktiv'}
"""

        QMessageBox.information(self, "RAG-Statistiken", info)
